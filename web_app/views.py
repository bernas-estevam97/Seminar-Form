from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.conf import settings
from django.http import HttpResponse, HttpResponseRedirect
from .forms import SeminarForm
from .models import SeminarFormModel, WordDocument
# from django.contrib.auth.models import User
# from django.contrib.auth import get_user_model
from docx import Document
import sqlite3
import os
from django.core.files.base import ContentFile
from docx.enum.style import WD_STYLE_TYPE
from django.template.loader import get_template
from xhtml2pdf import pisa
from datetime import date
from django.shortcuts import get_object_or_404
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.contrib import messages
from django.views.decorators.http import require_POST
import re
import requests

@login_required(login_url=settings.LOGIN_URL)
def home_page(request):
     main_page=True
     form = SeminarForm(request.POST)
     data = SeminarFormModel.objects.filter(user=request.user)
     return render(request, 'index.html', {'form': form, 'data': data, 'main_page': main_page})

@login_required(login_url=settings.LOGIN_URL)
# def seminar_form_req(request):
#      submitted=False
#      if request.method=='POST':
#           form = SeminarForm(request.POST)
#           if form.is_valid():
#                title = form.cleaned_data['seminar_title']
#                speaker = form.cleaned_data['seminar_speaker']
#                date = form.cleaned_data['seminar_date']

#                seminar_reg = SeminarFormModel.objects.create(
#                     user = request.user,
#                     seminar_title = title,
#                     seminar_speaker = speaker,
#                     seminar_date = date
#                )

#                seminar_reg.save()
#                # return render(request, "success.html", {'form': form, 'data': data})
#                return HttpResponseRedirect('/add-form?submitted=True')
#           # else:   
#           #      form = SeminarForm()
#           #      message = 'Something went wrong with the submission. Try again.'
#           #      return render(request, 'index.html', {'form': form, 'message': message})
#      else:
#           form = SeminarForm()
#           if 'submitted' in request.GET:
#                submitted = True
#      return render(request, 'index.html', {'form': form, 'submitted': submitted})

def seminar_form_req(request):
    if request.method == 'POST':
        form = SeminarForm(request.POST)
        recaptcha_token = request.POST.get('recaptcha_token')

        # Verify reCAPTCHA token with Google's API
        secret_key = settings.RECAPTCHA_PRIVATE_KEY  # Store your private key in settings.py
        recaptcha_url = 'https://www.google.com/recaptcha/api/siteverify'
        recaptcha_data = {
            'secret': secret_key,
            'response': recaptcha_token,
        }
        recaptcha_response = requests.post(recaptcha_url, data=recaptcha_data)
        recaptcha_result = recaptcha_response.json()

        # If reCAPTCHA verification fails, show an error message
        if not recaptcha_result.get('success'):
            messages.error(request, 'reCAPTCHA verification failed. Please try again.')
            return render(request, 'index.html', {'form': form})
        
        # Check the reCAPTCHA score
        score = recaptcha_result.get('score', 0)
        if score < settings.RECAPTCHA_REQUIRED_SCORE:
            messages.error(request, f"reCAPTCHA score of {score} is too low. Please try again.")
            return render(request, 'index.html', {'form': form})
        
        # If reCAPTCHA is successful, process the form
        if form.is_valid():
            title = form.cleaned_data['seminar_title']
            speaker = form.cleaned_data['seminar_speaker']
            date = form.cleaned_data['seminar_date']

            seminar_reg = SeminarFormModel.objects.create(
                user=request.user,
                seminar_title=title,
                seminar_speaker=speaker,
                seminar_date=date
            )
            seminar_reg.save()

            # Success message after successful form submission
            messages.success(request, 'Your seminar form was submitted successfully!')
            return redirect('home')  # Redirect to the same form page after success
        
        # If form is invalid, show an error message
        else:
            messages.error(request, 'There were errors with the form submission.')

    else:
        form = SeminarForm()

    return render(request, 'index.html', {'form': form})




def all_seminars(request):
     data = SeminarFormModel.objects.filter(user=request.user)
     return render(request, 'form-list.html',  {'seminars': data})


# Generate report in word (it doesn't work as intended - python-docx isn't great :)

# def generate_report(request):
#      data = SeminarFormModel.objects.filter(user=request.user)
#      report_generated=False
#      user=request.user
#      # Step 1: Connect to SQLite database and retrieve data
#      db_path = r"C:\Users\berna\Desktop\Coding projects\Python Projects\seminar_form\db.sqlite3"  # Path to your SQLite database
#      query = "SELECT seminar_title, seminar_speaker, seminar_date FROM web_app_seminarformmodel WHERE user_id="+str(user.id)  # Replace with your query
#      conn = sqlite3.connect(db_path)
#      cursor = conn.cursor()
#      cursor.execute(query)

#      # Fetch all rows and column names
#      rows = cursor.fetchall()
#      columns = [description[0] for description in cursor.description]

#      # Close the database connection
#      conn.close()
#      doc_path=r'C:\Users\berna\Desktop\Coding projects\Python Projects\seminar_form\media\seminar_declaration.docx'
#      document = Document(doc_path)
#      styles = document.styles
#      styles.add_style('List Bullet', WD_STYLE_TYPE.PARAGRAPH, builtin=True)
#      num_rows = len(rows) + 1  # +1 for the header row
#      num_cols = len(columns)
#      document.add_page_break()
     
#      table = document.add_table(rows=num_rows, cols=num_cols)
     
#      # Add the header row
#      header_row = table.rows[0]
#      header_row.cells[0].text = 'Seminar Title'
#      header_row.cells[1].text = 'Seminar Speaker'
#      header_row.cells[2].text = 'Seminar Date'

#      # Add the data rows
#      for i, row_data in enumerate(rows, start=1):  # Start from 1 because 0 is the header
#           row = table.rows[i]
#           for j, cell_data in enumerate(row_data):
#                row.cells[j].text = str(cell_data)  # Convert to string if not already
#      table.style = 'Table Grid'
#      # Step 4: Save the updated document
#      # document.save(f'C:/Users/berna/Documents/Coding Projects/Python App/seminar_form/media/seminar_form_report_{user}.docx')
#      # if os.path.exists(f'C:/Users/berna/Documents/Coding Projects/Python App/seminar_form/media/seminar_form_report_{user}.docx'):
#      #      report_generated = True
#      from io import BytesIO
#      buffer = BytesIO()
#      document.save(buffer)
#      buffer.seek(0)

#      # Step 5: Save the file to the Django database
#      word_file = WordDocument(name=f"Generated Report {user}")
#      word_file.file.save(f"generated_report_{user}.docx", ContentFile(buffer.read()))
#      buffer.close()
#      return HttpResponse("Word document updated and saved to the database.")
#      # else:
#      #      report_generated = False
#      #      return render(request, 'form-list.html',  {'report_generated': report_generated, 'data': data})



@require_POST
@login_required
def delete_seminar(request, seminar_id):
    if request.headers.get("x-requested-with") == "XMLHttpRequest":
        try:
            seminar = SeminarFormModel.objects.get(id=seminar_id)
            seminar.delete()
            return JsonResponse({"success": True})
        except SeminarFormModel.DoesNotExist:
            return JsonResponse({"success": False, "error": "Seminar not found"}, status=404)
    return JsonResponse({"success": False, "error": "Invalid request"}, status=400)


@require_POST
@login_required
def delete_all_seminars(request):
    if request.headers.get("x-requested-with") == "XMLHttpRequest":
        SeminarFormModel.objects.all().delete()
        return JsonResponse({"success": True})
    return JsonResponse({"success": False}, status=400)

# If you want to get by title use this instead:
    #if request.method == "POST":
    #   seminar = get_object_or_404(SeminarFormModel, seminar_title=seminar_title)
    #   seminar.delete()
    #    return JsonResponse({"success": True})
    #return JsonResponse({"success": False}, status=400)

@login_required
def generate_report(request):
     data = SeminarFormModel.objects.filter(user=request.user)
     current_date = date.today()
     current_date_pt = f'{current_date.day}/{current_date.month}/{current_date.year}'
     username = request.user
     context = {
          'data': data,
          'current_date_pt': current_date_pt,
          'user': username
     }
     template = get_template('report_template_pt.html')
     html = template.render(context)
     response = HttpResponse(content_type='application/pdf')
     safe_firstname = re.sub(r'\W+', '_', request.user.first_name)
     safe_lastname = re.sub(r'\W+', '_', request.user.last_name)
     response['Content-Disposition'] = f'attachment; filename="report_{safe_firstname}_{safe_lastname}.pdf"'
     pisa_status = pisa.CreatePDF(html, dest=response)
     if len(data) != 0:
     
          # Check for errors
          if pisa_status.err:
               return HttpResponse('We had some errors <pre>' + html + '</pre>')
          return response 
     else:
#           return HttpResponse("""
#         <script>
#             alert('Your seminar list is empty! Make sure you have at least 1 seminar to generate a report');
#             window.location.href = '/';  // Redirect after alert
#         </script>
#     """)
          messages.error(request, "Your seminar list is empty! Make sure you have at least 1 seminar to generate a report")
          return redirect('/form-list')

# EXECEPTIONS

def error_403(request, exception):
     return render(request, 'accounts/blocked.html')


def error_404(request, exception):
     return render(request, 'accounts/not-found.html')

def error_500(request):
     return render(request, '500.html', status=500)
